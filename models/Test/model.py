import pandas as pd
from pysentimiento import create_analyzer
from matplotlib import pyplot as plt
from docx import Document
import numpy as np
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from matplotlib.backends.backend_agg import FigureCanvasAgg
from docx.shared import Inches
import inflection as inf



#para verificar

file_path = '/Users/mauriciodiaz/Desktop/scrap_16_junio_2023.csv'

#hacer diferentes funciones con los nombres

def inicia(filePath):
    # Leemos el excel
    try:
        df = pd.read_excel(file_path)
    except UnicodeDecodeError:
        print('Error ')

    #tomamos la columna de los reviews
    reviews = df['review_text']
    reviews.dropna(inplace=True)


    #En este metodo vamos a limpiar las reviews y ponerlas todas en ingles y resumirlas
    reviews = aplica(reviews)
    
    #conseguimos el nombre del restaurante
    nombre = obtenNombre(df)
    
    
    #creamos diccionario con valores de sentiment
    resp = {'NEU':0,'POS':0,'NEG':0}
    # juan = reviews.apply(verificaSentiment)
    # for sentiment in juan:
    #     print(sentiment)
    #     resp[sentiment] += 1
    # print(resp)

    # print(imprimeBarChart(resp))

    #regresa las palabras mas repetidas
    reviewsAux = reviews.apply(resumeReview)
    palRep = palabrasMasRepetidas(reviewsAux)

    
    altos = reviewsDeNivel(df,1)
    bajos = reviewsDeNivel(df,0)

    nivel = [altos, bajos]
    
    otro = []
    for serie in nivel:
        aux = aplica(serie)
        serie = aux.apply(resumeReview)
        otro.append(serie)
    
    
    
    #acaba lo de imprimir palabras negativas y positivass de reviews
    final = []
    for rep in otro:
        pal = palabrasMasRepetidas(rep)
        pal = pal.iloc[:30]
        final.append(pal)

    chart = imprimeBarChart(resp,notacionCamello(nombre))

    createDoc(chart,final,nombre)

   
    #print(reviews.iloc[9])

def notacionCamello(name):
    notCamello = inf.camelize(name, uppercase_first_letter=True)
    return notCamello
    

def obtenNombre(df):
    string = df.iloc[1,'name']
    limit = string.find('|')
    newString = string[:limit]
    newString = newString.rstrip('\n')
    return newString





def aplica(serie):
    serie.dropna(inplace=True)
    final = serie.apply(limpiaReviews)
    final = final.apply(lambda x: x.lower())
    return final
    

def reviewsDeNivel(df,num):
    if num == 1:
        filt = (df['review_rating'] > 3)
        reviewsFin = df.loc[filt,'review_text']
    else:
        filt = (df['review_rating'] < 4)
        reviewsFin = df.loc[filt,'review_text']

    return reviewsFin



    

# quita las stopwords de cada review para obtener las palabras clave
def resumeReview(text):
    stop_words = set(stopwords.words('english'))
    split = text.split(' ')

    filter = []
    for word in split:
        if word not in stop_words:
            filter.append(word)
    final = ''
    for word in filter:
        final += word + ' '

    return final
    
   

#metodo convierte serie en una cadena larga de palabras para luego convertirla en otra serie 
#   de solamente palabras y sacar la cuenta de cada una
def palabrasMasRepetidas(serie):
    cadenaLarga = ' '.join(serie)
    listaDePalabras = cadenaLarga.split()
    seriePal = pd.Series(listaDePalabras)
    return seriePal.value_counts()


#metodo que limpia cada review dejando solo texto en ingles
def limpiaReviews(line):
    if line.__contains__('(Translated by Google)'):
        inicio = '(Translated by Google)'
        index = line.find('(Original)')
        line = line[len(inicio):index]
        line = line.rstrip('\n')
    
    return line


#me guarda un Bar Chart en una variable para poder imprimirla
def imprimeBarChart(dic,name):
    
    sent = list(dic.keys())
    value = list(dic.values())
    

    plt.bar(sent,value,color='red',width=0.4)
    plt.xlabel('Sentimiento')
    plt.ylabel('# de personas')
    plt.title('# de personas con cada sentimiento')

    plt.bar(sent, value)

    # Create a FigureCanvasAgg instance and render the figure onto it
    fig = plt.gcf()
    canvas = FigureCanvasAgg(fig)
    canvas.draw()

    # Save the rendered figure as a variable (in this example, as a BytesIO object)
    import io
    buffer = io.BytesIO()
    canvas.print_png(buffer)
    chart_data = buffer.getvalue()

    # You can now use the 'chart_data' variable, for example, save it to a file
    with open(f'{name}.png', 'wb') as f:
        f.write(chart_data)
    
    return f'{name}.png'



def verificaSentiment(line):
    analyzer = create_analyzer(task='sentiment',lang='en')
    answer = analyzer.predict(line)
    #regresa el si el texto es positivo negativo o neutral

    return answer.output

def createDoc(chart,list,name):
    doc = Document()
    doc.add_heading(name,level=1)
    doc.add_heading('Analysis',level=2)
    doc.add_paragraph('Palabras ocurrentes:\n# de estrellas > 3')
    doc.add_page_break()
    doc.add_paragraph(list[0])
    doc.add_section()
    doc.add_paragraph('Palabras ocurrentes:\n# de estrellas <= 3')
    doc.add_page_break()
    doc.add_paragraph(list[1])
    doc.add_paragraph('Sentiment Analysis')
    doc.add_picture(chart, width=Inches(6), height=Inches(4))

    # Save the Word document
    doc.save(f'Analysis{name}.docx')


    return doc










#inicia(file_path)





